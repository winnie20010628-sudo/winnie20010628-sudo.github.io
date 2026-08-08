from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


VOCABULARY = {
    "K1 (Ages 3–4)": {
        "Me and My Family": [
            ("I", "我", "I am happy."),
            ("you", "你", "You are my friend."),
            ("Mum / Mommy", "媽媽", "Mum loves me."),
            ("Dad / Daddy", "爸爸", "Dad is tall."),
            ("baby", "嬰兒", "The baby is sleeping."),
            ("boy", "男孩", "The boy can run."),
            ("girl", "女孩", "The girl has a doll."),
            ("friend", "朋友", "My friend is kind."),
        ],
        "Classroom and Toys": [
            ("book", "書", "This is my book."),
            ("bag", "書包", "My bag is blue."),
            ("chair", "椅子", "Sit on the chair."),
            ("table", "桌子", "The table is big."),
            ("pencil", "鉛筆", "I have a pencil."),
            ("ball", "球", "Throw the ball."),
            ("toy", "玩具", "It is a new toy."),
            ("doll", "洋娃娃", "The doll is pretty."),
        ],
        "Colours, Numbers and Shapes": [
            ("red", "紅色", "The apple is red."),
            ("blue", "藍色", "The sky is blue."),
            ("yellow", "黃色", "The sun is yellow."),
            ("green", "綠色", "The leaf is green."),
            ("one", "一", "I have one ball."),
            ("two", "二", "I see two cats."),
            ("circle", "圓形", "It is a circle."),
            ("star", "星形", "Draw a star."),
        ],
        "Everyday Words": [
            ("hello", "你好", "Hello, teacher!"),
            ("goodbye", "再見", "Goodbye, Mum!"),
            ("please", "請", "Please help me."),
            ("thank you", "謝謝", "Thank you, Dad."),
            ("yes", "是", "Yes, please."),
            ("no", "不", "No, thank you."),
            ("happy", "快樂的", "I am happy."),
            ("sad", "傷心的", "The baby is sad."),
        ],
    },
    "K2 (Ages 4–5)": {
        "People and Places": [
            ("teacher", "老師", "My teacher is nice."),
            ("doctor", "醫生", "The doctor helps me."),
            ("police officer", "警察", "The police officer keeps us safe."),
            ("school", "學校", "I go to school."),
            ("home", "家", "I am at home."),
            ("park", "公園", "We play in the park."),
            ("shop", "商店", "The shop has food."),
            ("hospital", "醫院", "The hospital is near."),
        ],
        "Food and Healthy Habits": [
            ("apple", "蘋果", "I eat an apple."),
            ("banana", "香蕉", "The banana is yellow."),
            ("rice", "米飯", "I like rice."),
            ("milk", "牛奶", "Drink some milk."),
            ("water", "水", "Water is good for me."),
            ("breakfast", "早餐", "I eat breakfast."),
            ("wash", "洗", "Wash your hands."),
            ("teeth", "牙齒", "Brush your teeth."),
        ],
        "Animals and Nature": [
            ("cat", "貓", "The cat says meow."),
            ("dog", "狗", "The dog is brown."),
            ("bird", "鳥", "The bird can fly."),
            ("fish", "魚", "The fish can swim."),
            ("tree", "樹", "The tree is tall."),
            ("flower", "花", "The flower is pink."),
            ("sun", "太陽", "The sun is hot."),
            ("rain", "雨", "I hear the rain."),
        ],
        "Actions and Describing Words": [
            ("run", "跑", "I can run fast."),
            ("jump", "跳", "Jump up high."),
            ("walk", "走路", "We walk to school."),
            ("eat", "吃", "Eat your lunch."),
            ("drink", "喝", "Drink water."),
            ("big", "大的", "The elephant is big."),
            ("small", "小的", "The ant is small."),
            ("fast", "快的", "The car is fast."),
        ],
    },
    "K3 (Ages 5–6)": {
        "School Readiness": [
            ("alphabet", "字母表", "The alphabet has 26 letters."),
            ("letter", "字母", "A is a letter."),
            ("word", "單字", "Read the word."),
            ("sentence", "句子", "Write a sentence."),
            ("read", "閱讀", "I can read a book."),
            ("write", "書寫", "Write your name."),
            ("listen", "聆聽", "Listen to the story."),
            ("answer", "回答", "Answer the question."),
        ],
        "Time, Weather and Seasons": [
            ("morning", "早上", "Good morning!"),
            ("afternoon", "下午", "See you this afternoon."),
            ("today", "今天", "Today is Monday."),
            ("tomorrow", "明天", "Tomorrow is a holiday."),
            ("sunny", "晴朗的", "It is sunny today."),
            ("cloudy", "多雲的", "It is cloudy now."),
            ("winter", "冬天", "Winter is cool."),
            ("summer", "夏天", "Summer is hot."),
        ],
        "Community and Hong Kong": [
            ("Hong Kong", "香港", "I live in Hong Kong."),
            ("MTR", "港鐵", "We take the MTR."),
            ("bus", "巴士", "The bus stops here."),
            ("cross", "過馬路", "Cross at the green light."),
            ("traffic light", "交通燈", "Wait at the traffic light."),
            ("library", "圖書館", "We borrow books at the library."),
            ("recycle", "回收", "Recycle paper and cans."),
            ("clean", "乾淨的", "Keep our city clean."),
        ],
        "Feelings and Good Manners": [
            ("excited", "興奮的", "I am excited today."),
            ("worried", "擔心的", "Do not be worried."),
            ("proud", "自豪的", "I am proud of you."),
            ("share", "分享", "Share your toys."),
            ("kind", "友善的", "Be kind to friends."),
            ("sorry", "對不起", "I am sorry."),
            ("excuse me", "勞駕／對不起", "Excuse me, please."),
            ("help", "幫助", "Can you help me?"),
        ],
    },
}


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Hong Kong Kindergarten English Vocabulary • K1–K3")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)
add_footer(section)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Hong Kong K1–K3 English Vocabulary List")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(39, 91, 157)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Essential words for kindergarten learners • English | 中文 | Example sentence")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(90, 90, 90)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Tip: Practise 5–8 words at a time through songs, picture cards, stories, and daily routines.")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(80, 120, 80)

for level_index, (level, topics) in enumerate(VOCABULARY.items()):
    doc.add_heading(level, level=1)
    for topic, words in topics.items():
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(3)
        run = heading.add_run(topic)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(39, 91, 157)

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        headers = ["English word", "中文意思", "Example"]
        for cell, text in zip(table.rows[0].cells, headers):
            set_cell_text(cell, text, bold=True, color=(255, 255, 255))
            shade(cell, "275B9D")

        for row_number, (word, chinese, example) in enumerate(words):
            cells = table.add_row().cells
            for cell, text in zip(cells, [word, chinese, example]):
                set_cell_text(cell, text)
                if row_number % 2 == 0:
                    shade(cell, "EAF2F8")

        table.columns[0].width = Inches(1.45)
        table.columns[1].width = Inches(1.45)
        table.columns[2].width = Inches(3.75)

    if level_index < len(VOCABULARY) - 1:
        doc.add_page_break()

doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("Well done! Keep listening, speaking, reading, and playing in English.")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(39, 91, 157)

doc.core_properties.title = "Hong Kong K1–K3 English Vocabulary List"
doc.core_properties.subject = "Essential English vocabulary for kindergarten learners"
doc.core_properties.author = "Cursor"
doc.save("Hong_Kong_K1-K3_English_Vocabulary_List.docx")
